#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
增强版 Markdown 转 Word 文档转换器

集成 pandoc 作为主要转换引擎，支持：
1. 页码
2. 页眉页脚
3. 自动目录
4. 分页符
5. 图表标题格式化

使用方法：
    python scripts/enhanced_md_to_docx.py --input paper.md --output paper.docx --template template.docx
"""

import re
import os
import sys
import subprocess
import argparse
from pathlib import Path
from datetime import datetime
from typing import Optional, Tuple, List

from terminal_encoding import subprocess_text_kwargs

# 检查 python-docx 是否可用
try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsmap
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("[警告] python-docx 未安装，后处理功能不可用")
    print("安装命令: pip install python-docx")


def check_pandoc_installed() -> bool:
    """检查 pandoc 是否已安装"""
    try:
        result = subprocess.run(
            ['pandoc', '--version'],
            capture_output=True,
            timeout=10,
            **subprocess_text_kwargs()
        )
        return result.returncode == 0
    except (FileNotFoundError, subprocess.TimeoutExpired):
        return False


def clean_markdown_content(content: str) -> str:
    """
    清理 Markdown 内容中的无意义字符

    Args:
        content: 原始 Markdown 内容

    Returns:
        清理后的内容
    """
    # 移除连续多个 ---（3个以上的水平线）
    content = re.sub(r'\n-{3,}\n', '\n\n', content)
    content = re.sub(r'\n-{3,}\s*\n', '\n\n', content)

    # 移除开头的 ---
    content = re.sub(r'^-{3,}\s*\n', '', content)

    # 移除结尾的 ---
    content = re.sub(r'\n\s*-{3,}$', '', content)

    # 移除多余的星号分隔线
    content = re.sub(r'\n\*{3,}\n', '\n\n', content)

    # 压缩多个空行为最多两个
    content = re.sub(r'\n{4,}', '\n\n\n', content)

    # 移除行尾空白
    content = re.sub(r'[ \t]+\n', '\n', content)

    return content


def convert_with_pandoc(
    input_path: str,
    output_path: str,
    template_path: Optional[str] = None,
    toc: bool = True,
    toc_depth: int = 3
) -> Tuple[bool, str]:
    """
    使用 pandoc 进行转换

    Args:
        input_path: 输入 Markdown 文件路径
        output_path: 输出 Word 文件路径
        template_path: Word 模板文件路径
        toc: 是否生成目录
        toc_depth: 目录深度

    Returns:
        (success, message)
    """
    if not check_pandoc_installed():
        return False, "pandoc 未安装，请先安装：winget install pandoc 或访问 https://pandoc.org/installing.html"

    try:
        # 构建 pandoc 命令
        cmd = [
            'pandoc',
            input_path,
            '-o', output_path,
            '-f', 'markdown',
            '-t', 'docx',
        ]

        # 添加模板
        if template_path and Path(template_path).exists():
            cmd.extend(['--reference-doc', template_path])

        # 添加目录
        if toc:
            cmd.extend(['--toc', '--toc-depth', str(toc_depth)])

        # 执行转换
        result = subprocess.run(
            cmd,
            capture_output=True,
            timeout=120,
            **subprocess_text_kwargs()
        )

        if result.returncode == 0:
            return True, f"转换成功：{output_path}"
        else:
            return False, f"pandoc 转换失败：{result.stderr}"

    except subprocess.TimeoutExpired:
        return False, "pandoc 转换超时"
    except Exception as e:
        return False, f"转换异常：{str(e)}"


def add_page_numbers(doc_path: str, start_from: int = 1) -> bool:
    """
    为 Word 文档添加页码

    Args:
        doc_path: Word 文档路径
        start_from: 起始页码

    Returns:
        是否成功
    """
    if not DOCX_AVAILABLE:
        return False

    try:
        doc = Document(doc_path)

        for section in doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False

            # 清空现有页脚内容
            for para in footer.paragraphs:
                para.clear()

            # 添加页码
            para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 创建页码域
            run = para.add_run()
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')

            instrText = OxmlElement('w:instrText')
            instrText.text = f"PAGE \\* MERGEFORMAT"

            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')

            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)

            # 设置字体
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)

        doc.save(doc_path)
        return True

    except Exception as e:
        print(f"[警告] 添加页码失败：{e}")
        return False


def add_headers(doc_path: str, header_text: str = "") -> bool:
    """
    为 Word 文档添加页眉

    Args:
        doc_path: Word 文档路径
        header_text: 页眉文本

    Returns:
        是否成功
    """
    if not DOCX_AVAILABLE:
        return False

    try:
        doc = Document(doc_path)

        for section in doc.sections:
            header = section.header
            header.is_linked_to_previous = False

            # 清空现有页眉内容
            for para in header.paragraphs:
                para.clear()

            # 添加页眉
            para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run = para.add_run(header_text)
            run.font.name = '宋体'
            run.font.size = Pt(10)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            # 添加下划线
            para_format = para.paragraph_format
            para_format.space_after = Pt(6)

        doc.save(doc_path)
        return True

    except Exception as e:
        print(f"[警告] 添加页眉失败：{e}")
        return False


def set_page_margins(doc_path: str, top: float = 2.54, bottom: float = 2.54,
                     left: float = 3.17, right: float = 3.17) -> bool:
    """
    设置页面边距（单位：厘米）

    Args:
        doc_path: Word 文档路径
        top, bottom, left, right: 边距值（厘米）

    Returns:
        是否成功
    """
    if not DOCX_AVAILABLE:
        return False

    try:
        doc = Document(doc_path)

        for section in doc.sections:
            section.top_margin = Cm(top)
            section.bottom_margin = Cm(bottom)
            section.left_margin = Cm(left)
            section.right_margin = Cm(right)

        doc.save(doc_path)
        return True

    except Exception as e:
        print(f"[警告] 设置页边距失败：{e}")
        return False


def add_page_break_before_section(doc_path: str, section_titles: List[str]) -> bool:
    """
    在指定章节前添加分页符

    Args:
        doc_path: Word 文档路径
        section_titles: 需要在其前添加分页符的章节标题列表

    Returns:
        是否成功
    """
    if not DOCX_AVAILABLE:
        return False

    try:
        doc = Document(doc_path)

        for para in doc.paragraphs:
            # 检查是否为目标章节标题
            for title in section_titles:
                if title in para.text and para.style.name.startswith('Heading'):
                    # 在段落前插入分页符
                    para.paragraph_format.page_break_before = True
                    break

        doc.save(doc_path)
        return True

    except Exception as e:
        print(f"[警告] 添加分页符失败：{e}")
        return False


def format_figure_captions(doc_path: str) -> bool:
    """
    格式化图表标题

    Args:
        doc_path: Word 文档路径

    Returns:
        是否成功
    """
    if not DOCX_AVAILABLE:
        return False

    try:
        doc = Document(doc_path)

        # 图标题模式：图X-X
        figure_pattern = re.compile(r'^图\d+-\d+')
        # 表标题模式：表X-X
        table_pattern = re.compile(r'^表\d+-\d+')

        for para in doc.paragraphs:
            text = para.text.strip()

            if figure_pattern.match(text):
                # 图标题：居中、宋体、五号
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            elif table_pattern.match(text):
                # 表标题：居中、宋体、五号
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        doc.save(doc_path)
        return True

    except Exception as e:
        print(f"[警告] 格式化图表标题失败：{e}")
        return False


def convert_md_to_docx(
    input_path: str,
    output_path: str,
    template_path: Optional[str] = None,
    header_text: str = "",
    page_number_start: int = 1,
    toc: bool = True,
    clean_content: bool = True
) -> Tuple[bool, str]:
    """
    完整的 Markdown 转 Word 流程

    Args:
        input_path: 输入 Markdown 文件路径
        output_path: 输出 Word 文件路径
        template_path: Word 模板文件路径
        header_text: 页眉文本
        page_number_start: 起始页码
        toc: 是否生成目录
        clean_content: 是否清理内容

    Returns:
        (success, message)
    """
    print(f"[信息] 开始转换: {input_path}")

    # Step 1: 读取并清理内容
    if clean_content:
        with open(input_path, 'r', encoding='utf-8') as f:
            content = f.read()

        content = clean_markdown_content(content)

        # 创建临时清理后的文件
        temp_input = Path(input_path).with_suffix('.cleaned.md')
        with open(temp_input, 'w', encoding='utf-8') as f:
            f.write(content)

        actual_input = str(temp_input)
    else:
        actual_input = input_path

    # Step 2: 使用 pandoc 转换
    success, message = convert_with_pandoc(
        actual_input,
        output_path,
        template_path,
        toc=toc
    )

    if not success:
        # 清理临时文件
        if clean_content and Path(actual_input).exists():
            Path(actual_input).unlink()
        return False, message

    print(f"[成功] pandoc 转换完成")

    # Step 3: 后处理（添加页码、页眉等）
    if DOCX_AVAILABLE:
        # 添加页码
        if add_page_numbers(output_path, page_number_start):
            print(f"[成功] 已添加页码")

        # 添加页眉
        if header_text and add_headers(output_path, header_text):
            print(f"[成功] 已添加页眉")

        # 设置页边距
        if set_page_margins(output_path):
            print(f"[成功] 已设置页边距")

        # 格式化图表标题
        if format_figure_captions(output_path):
            print(f"[成功] 已格式化图表标题")

    # 清理临时文件
    if clean_content and Path(actual_input).exists():
        Path(actual_input).unlink()

    return True, f"Word 文档已保存到: {output_path}"


def main():
    """main"""
    parser = argparse.ArgumentParser(description='增强版 Markdown 转 Word 文档转换器')
    parser.add_argument('--input', '-i', required=True, help='输入 Markdown 文件路径')
    parser.add_argument('--output', '-o', help='输出 Word 文件路径（默认：同名 .docx）')
    parser.add_argument('--template', '-t', help='Word 模板文件路径')
    parser.add_argument('--header', '-H', default='', help='页眉文本')
    parser.add_argument('--page-start', '-p', type=int, default=1, help='起始页码')
    parser.add_argument('--no-toc', action='store_true', help='不生成目录')
    parser.add_argument('--no-clean', action='store_true', help='不清理内容')

    args = parser.parse_args()

    # 设置输出路径
    input_path = Path(args.input)
    output_path = args.output or str(input_path.with_suffix('.docx'))

    # 执行转换
    success, message = convert_md_to_docx(
        input_path=str(input_path),
        output_path=output_path,
        template_path=args.template,
        header_text=args.header,
        page_number_start=args.page_start,
        toc=not args.no_toc,
        clean_content=not args.no_clean
    )

    if success:
        print(f"\n[完成] {message}")
    else:
        print(f"\n[失败] {message}")
        sys.exit(1)


if __name__ == '__main__':
    main()

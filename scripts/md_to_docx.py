#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Markdown 转 Word 文档转换器
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_chinese_font(run, font_name='宋体', font_size=12):
    """设置中文字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def create_thesis_document():
    """创建论文Word文档"""
    doc = Document()

    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    return doc


def add_heading(doc, text, level):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        set_chinese_font(run, '黑体', 16 - level * 2)
    return heading


def add_paragraph(doc, text, first_line_indent=True):
    """添加段落"""
    para = doc.add_paragraph()
    if first_line_indent:
        para.paragraph_format.first_line_indent = Cm(0.74)
    para.paragraph_format.line_spacing = 1.5

    run = para.add_run(text)
    set_chinese_font(run, '宋体', 12)
    return para


def parse_markdown_line(line):
    """解析Markdown行"""
    # 标题
    if line.startswith('# '):
        return ('h1', line[2:])
    elif line.startswith('## '):
        return ('h2', line[3:])
    elif line.startswith('### '):
        return ('h3', line[4:])
    elif line.startswith('#### '):
        return ('h4', line[5:])

    # 列表
    if line.startswith('- ') or line.startswith('* '):
        return ('list', line[2:])
    if re.match(r'^\d+\.\s', line):
        return ('list', re.sub(r'^\d+\.\s', '', line))

    # 表格
    if line.startswith('|'):
        return ('table', line)

    # 代码块
    if line.startswith('```'):
        return ('code', line)

    # 普通段落
    if line.strip():
        return ('para', line)

    return ('empty', '')


def process_inline_formatting(doc, para, text):
    """处理行内格式（加粗、斜体、代码）"""
    # 简化处理，直接添加文本
    run = para.add_run(text)
    set_chinese_font(run, '宋体', 12)
    return para


def convert_md_to_docx(input_path, output_path):
    """将Markdown转换为Word文档"""
    print(f"正在读取: {input_path}")

    with open(input_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')

    doc = create_thesis_document()

    in_code_block = False
    code_buffer = []
    in_table = False
    table_rows = []

    for line in lines:
        parsed = parse_markdown_line(line)
        line_type, line_content = parsed

        # 代码块处理
        if line_type == 'code':
            if in_code_block:
                # 结束代码块
                if code_buffer:
                    code_text = '\n'.join(code_buffer)
                    para = doc.add_paragraph()
                    para.paragraph_format.left_indent = Cm(1)
                    run = para.add_run(code_text)
                    set_chinese_font(run, 'Consolas', 10)
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_buffer.append(line)
            continue

        # 表格处理
        if line_type == 'table':
            if not in_table:
                in_table = True
                table_rows = []
            # 跳过分隔行
            if not re.match(r'^\|[\s\-:]+\|', line):
                cells = [c.strip() for c in line.split('|')[1:-1]]
                if cells:
                    table_rows.append(cells)
            continue
        else:
            if in_table and table_rows:
                # 创建表格
                if len(table_rows) > 0:
                    table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                    table.style = 'Table Grid'
                    for i, row_data in enumerate(table_rows):
                        for j, cell_text in enumerate(row_data):
                            if j < len(table.rows[i].cells):
                                cell = table.rows[i].cells[j]
                                cell.text = cell_text
                table_rows = []
                in_table = False

        # 标题处理
        if line_type == 'h1':
            # 居中标题
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(line_content)
            set_chinese_font(run, '黑体', 18)
            run.bold = True
        elif line_type == 'h2':
            add_heading(doc, line_content, 1)
        elif line_type == 'h3':
            add_heading(doc, line_content, 2)
        elif line_type == 'h4':
            add_heading(doc, line_content, 3)
        elif line_type == 'list':
            para = doc.add_paragraph(style='List Bullet')
            run = para.add_run(line_content)
            set_chinese_font(run, '宋体', 12)
        elif line_type == 'para':
            # 处理特殊标记
            text = line_content
            # 处理上标引用
            text = re.sub(r'<sup>\[(\d+)\]</sup>', r'[\1]', text)
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # 移除加粗标记
            text = re.sub(r'\*(.+?)\*', r'\1', text)  # 移除斜体标记
            text = re.sub(r'`(.+?)`', r'\1', text)  # 移除代码标记

            add_paragraph(doc, text)
        elif line_type == 'empty':
            pass  # 跳过空行

    # 处理最后的表格
    if in_table and table_rows:
        table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        table.style = 'Table Grid'
        for i, row_data in enumerate(table_rows):
            for j, cell_text in enumerate(row_data):
                if j < len(table.rows[i].cells):
                    table.rows[i].cells[j].text = cell_text

    # 保存文档
    doc.save(output_path)
    print(f"Word文档已保存到: {output_path}")
    return output_path


if __name__ == '__main__':
    import sys
    input_file = sys.argv[1] if len(sys.argv) > 1 else "workspace/final/论文终稿.md"
    output_file = sys.argv[2] if len(sys.argv) > 2 else "workspace/final/论文终稿.docx"

    convert_md_to_docx(input_file, output_file)
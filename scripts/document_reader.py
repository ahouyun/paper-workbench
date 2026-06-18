#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
文档读取器（Document Reader）

支持读取 PDF 和 Word 文档，提取文本内容和文档结构。

使用方法：
    python scripts/document_reader.py --input paper.pdf
    python scripts/document_reader.py --input paper.docx
    python scripts/document_reader.py --input paper.pdf --structure
"""

import re
import argparse
from pathlib import Path
from typing import Optional, Dict, List


class DocumentReader:
    """文档读取器"""

    def read_pdf(self, path: str) -> str:
        """
        读取 PDF 文件

        依赖: PyMuPDF (fitz)

        Args:
            path: PDF 文件路径

        Returns:
            提取的文本内容
        """
        try:
            import fitz  # PyMuPDF
        except ImportError:
            print("[错误] PyMuPDF 未安装，PDF 读取不可用")
            print("安装命令: pip install PyMuPDF")
            return ""

        try:
            doc = fitz.open(path)
            text_parts = []
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text("text")
                if text.strip():
                    text_parts.append(text.strip())
            doc.close()
            return "\n\n".join(text_parts)
        except Exception as e:
            print(f"[错误] PDF 读取失败: {e}")
            return ""

    def read_docx(self, path: str) -> str:
        """
        读取 Word 文档

        依赖: python-docx

        Args:
            path: Word 文件路径

        Returns:
            提取的文本内容
        """
        try:
            from docx import Document
        except ImportError:
            print("[错误] python-docx 未安装，Word 读取不可用")
            print("安装命令: pip install python-docx")
            return ""

        try:
            doc = Document(path)
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())

            # 读取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    text_parts.append(row_text)

            return "\n\n".join(text_parts)
        except Exception as e:
            print(f"[错误] Word 文档读取失败: {e}")
            return ""

    def read_auto(self, path: str) -> str:
        """自动识别文件类型并读取"""
        suffix = Path(path).suffix.lower()
        if suffix == '.pdf':
            return self.read_pdf(path)
        elif suffix in ('.docx', '.doc'):
            return self.read_docx(path)
        else:
            print(f"[错误] 不支持的文件格式: {suffix}")
            return ""

    def extract_structure(self, path: str) -> Dict:
        """
        提取文档结构（标题层级）

        Args:
            path: 文件路径

        Returns:
            结构字典，包含标题层级信息
        """
        suffix = Path(path).suffix.lower()
        structure = {
            "file": str(path),
            "type": suffix,
            "headings": [],
            "total_paragraphs": 0,
            "total_tables": 0
        }

        if suffix == '.pdf':
            # PDF 结构提取（基于文本模式匹配）
            text = self.read_pdf(path)
            lines = text.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                # 匹配类似 "第X章" 的标题
                if re.match(r'^第\d+章\s', line):
                    structure["headings"].append({"level": 1, "text": line})
                elif re.match(r'^\d+\.\d+\s', line):
                    structure["headings"].append({"level": 2, "text": line})
                elif re.match(r'^\d+\.\d+\.\d+\s', line):
                    structure["headings"].append({"level": 3, "text": line})

        elif suffix in ('.docx', '.doc'):
            try:
                from docx import Document
                doc = Document(path)
                structure["total_paragraphs"] = len(doc.paragraphs)
                structure["total_tables"] = len(doc.tables)

                for para in doc.paragraphs:
                    if para.style.name.startswith('Heading'):
                        level = int(para.style.name.replace('Heading ', '').replace('Heading', '1'))
                        structure["headings"].append({
                            "level": level,
                            "text": para.text.strip()
                        })
            except Exception as e:
                print(f"[错误] 文档结构提取失败: {e}")

        return structure


def main():
    """main"""
    parser = argparse.ArgumentParser(description="文档读取器")
    parser.add_argument("--input", "-i", required=True, help="输入文件路径（PDF/Word）")
    parser.add_argument("--structure", "-s", action="store_true", help="提取文档结构")
    parser.add_argument("--output", "-o", help="输出文件路径（可选，默认控制台输出）")

    args = parser.parse_args()
    reader = DocumentReader()

    if not Path(args.input).exists():
        print(f"[错误] 文件不存在: {args.input}")
        return

    if args.structure:
        result = reader.extract_structure(args.input)
        output_lines = [
            f"文件: {result['file']}",
            f"类型: {result['type']}",
            f"段落数: {result['total_paragraphs']}",
            f"表格数: {result['total_tables']}",
            "",
            "标题结构:",
        ]
        for h in result["headings"]:
            indent = "  " * h["level"]
            output_lines.append(f"{indent}[L{h['level']}] {h['text']}")
        output = "\n".join(output_lines)
    else:
        output = reader.read_auto(args.input)

    if args.output:
        Path(args.output).write_text(output, encoding='utf-8')
        print(f"[成功] 内容已保存到: {args.output}")
    else:
        print(output)


if __name__ == "__main__":
    main()

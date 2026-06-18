import re
from pathlib import Path
from typing import Tuple

from .config import FORMAT_CONFIG
from .markdown import clean_markdown_content, parse_markdown, strip_doi_links
from .preflight import preflight_validate_images

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("[警告] python-docx 未安装，Word 导出功能不可用")
    print("安装命令: pip install python-docx")

try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
    print("[警告] Pillow 未安装，图片尺寸自动缩放功能不可用")
    print("安装命令: pip install Pillow")


def set_chinese_font(run, font_name: str = '宋体', font_size: int = 12, bold: bool = False):
    """设置中文字体（支持加粗）"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold

    # 设置中文字体（East Asia）
    r = run._element
    rPr = r.rPr
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)

    # 设置 East Asia 字体
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)

    rFonts.set(qn('w:eastAsia'), font_name)

    # 中文字体加粗需要额外设置 bCs 属性
    if bold:
        # 设置西文加粗
        b_elem = rPr.find(qn('w:b'))
        if b_elem is None:
            b_elem = OxmlElement('w:b')
            rPr.append(b_elem)

        # 设置东亚文字加粗（关键！）
        bCs_elem = rPr.find(qn('w:bCs'))
        if bCs_elem is None:
            bCs_elem = OxmlElement('w:bCs')
            rPr.append(bCs_elem)


def set_paragraph_format(para, first_line_indent: bool = True, line_spacing: float = 1.5):
    """设置段落格式"""
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if first_line_indent:
        para.paragraph_format.first_line_indent = Cm(0.74)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(0)


def create_thesis_document() -> 'Document':
    """创建论文 Word 文档（预设格式）"""
    doc = Document()

    # 设置页面边距（符合中国学术论文标准）
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    return doc


def add_title(doc, text: str):
    """添加论文标题（居中、黑体、三号）"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    set_chinese_font(run, '黑体', 16, bold=True)
    para.paragraph_format.space_after = Pt(12)


def add_heading(doc, text: str, level: int):
    """添加章节标题"""
    para = doc.add_paragraph()

    # 根据级别设置格式
    if level == 1:  # 一级标题（章标题）
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        set_chinese_font(run, '黑体', 14, bold=True)
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)
    elif level == 2:  # 二级标题（节标题）
        run = para.add_run(text)
        set_chinese_font(run, '黑体', 12, bold=True)
        para.paragraph_format.first_line_indent = Cm(0)
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(3)
    elif level == 3:  # 三级标题
        run = para.add_run(text)
        set_chinese_font(run, '黑体', 12, bold=True)
        para.paragraph_format.first_line_indent = Cm(0)
        para.paragraph_format.space_before = Pt(3)
        para.paragraph_format.space_after = Pt(3)
    else:
        run = para.add_run(text)
        set_chinese_font(run, '黑体', 12, bold=True)

    return para


def add_paragraph(doc, text: str, first_line_indent: bool = True):
    """添加正文段落"""
    para = doc.add_paragraph()
    set_paragraph_format(para, first_line_indent)

    # 处理行内格式
    process_inline_formatting(para, text)

    return para


def process_inline_formatting(para, text: str):
    """处理行内格式（加粗、斜体、代码、上标）"""
    # 处理上标引用 <sup>[1]</sup> -> [1]
    text = re.sub(r'<sup>\[(\d+)\]</sup>', r'[\1]', text)

    # 逐段处理 **加粗**、*斜体*、`代码`
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')
    last_end = 0

    for match in pattern.finditer(text):
        if match.start() > last_end:
            run = para.add_run(text[last_end:match.start()])
            set_chinese_font(run, '宋体', 12)

        if match.group(2):  # **加粗**
            run = para.add_run(match.group(2))
            set_chinese_font(run, '宋体', 12, bold=True)
        elif match.group(3):  # *斜体*
            run = para.add_run(match.group(3))
            set_chinese_font(run, '宋体', 12)
            run.font.italic = True
        elif match.group(4):  # `代码`
            run = para.add_run(match.group(4))
            run.font.name = 'Consolas'
            run.font.size = Pt(10.5)

        last_end = match.end()

    if last_end < len(text):
        run = para.add_run(text[last_end:])
        set_chinese_font(run, '宋体', 12)


def add_code_block(doc, code_lines: list, language: str = ''):
    """添加代码块"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(1)
    para.paragraph_format.first_line_indent = Cm(0)

    code_text = '\n'.join(code_lines)
    run = para.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)


def add_table(doc, rows: list, caption: str = ''):
    """
    添加三线表

    三线表规范：
    - 顶线：1.5pt 粗线（w:sz="12"）
    - 表头下线：0.75pt 细线（w:sz="6"）
    - 底线：1.5pt 粗线
    - 其余边框：无
    - 表头文字加粗居中

    Args:
        doc: Word 文档对象
        rows: 表格数据（第一行为表头）
        caption: 表名（在表上方显示，格式：表X.X  名称）
    """
    if not rows or len(rows) < 1:
        return

    # 添加表名（表名在表的上方）
    if caption:
        add_table_caption(doc, caption)

    # 创建表格
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 清除默认边框（先移除 Table Grid 样式的边框）
    table.style = 'Table Grid'

    # 填充内容
    table_font = FORMAT_CONFIG.get('table_font', '黑体')

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = table.rows[i].cells[j]
                cell.text = cell_text.strip()
                # 设置单元格字体
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        if i == 0:
                            # 表头加粗居中
                            set_chinese_font(run, table_font, 10, bold=True)
                        else:
                            set_chinese_font(run, table_font, 10)

    # 应用三线表边框
    _apply_three_line_borders(table)

    # 表格后添加空行
    doc.add_paragraph()


def _apply_three_line_borders(table):
    """
    应用三线表边框样式

    通过 OxmlElement 直接操作 OOXML，为每个单元格设置边框：
    - 第1行（表头）：顶线粗线 + 底线细线
    - 中间行：无边框
    - 最后1行：底线粗线
    """
    tbl = table._tbl

    # 定义边框尺寸（w:sz 单位为 1/8 磅）
    THICK_SZ = "12"   # 1.5pt = 12/8
    THIN_SZ = "6"     # 0.75pt = 6/8
    BORDER_COLOR = "000000"

    for i, row in enumerate(table.rows):
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.tcPr
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.insert(0, tcPr)

            # 移除已有边框
            existing = tcPr.find(qn('w:tcBorders'))
            if existing is not None:
                tcPr.remove(existing)

            tcBorders = OxmlElement('w:tcBorders')

            # 设置各边框
            for side in ['top', 'bottom', 'left', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{side}')
                border.set(qn('w:val'), 'none')
                border.set(qn('w:sz'), '0')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), BORDER_COLOR)

                if i == 0 and side == 'top':
                    # 顶线：粗线
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), THICK_SZ)
                elif i == 0 and side == 'bottom':
                    # 表头下线：细线
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), THIN_SZ)
                elif i == len(table.rows) - 1 and side == 'bottom':
                    # 底线：粗线
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), THICK_SZ)
                elif side in ('left', 'right'):
                    # 左右边框：无
                    pass

                tcBorders.append(border)

            tcPr.append(tcBorders)


def add_table_caption(doc, caption: str):
    """
    添加表名（表名在表的上方）

    格式：居中、宋体、五号（10.5pt）

    Args:
        doc: Word 文档对象
        caption: 表名文本，如 "表4.1  用户信息表结构"
    """
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.first_line_indent = Cm(0)
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(3)

    run = para.add_run(caption)
    set_chinese_font(run, '宋体', 10.5)


def add_list_item(doc, text: str, ordered: bool = False):
    """添加列表项"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.74)
    para.paragraph_format.first_line_indent = Cm(0)

    if ordered:
        run = para.add_run(text)
    else:
        run = para.add_run(f"• {text}")

    set_chinese_font(run, '宋体', 12)


def calculate_image_size(image_path: str, max_width_cm: float = 14.0, max_height_cm: float = 12.0) -> Tuple[float, float]:
    """
    计算图片在 Word 文档中的合适尺寸

    Args:
        image_path: 图片路径
        max_width_cm: 最大宽度（厘米），A4 纸有效宽度约 14cm
        max_height_cm: 最大高度（厘米），避免图片过长

    Returns:
        (width_cm, height_cm): 计算后的尺寸
    """
    if not PIL_AVAILABLE:
        # Pillow 不可用时使用默认宽度
        return max_width_cm, None

    try:
        with Image.open(image_path) as img:
            # 获取原始尺寸（像素）
            orig_width_px, orig_height_px = img.size

            # 假设 Word 文档 DPI 为 96（标准屏幕 DPI）
            # 1 英寸 = 2.54 厘米，1 厘米 ≈ 37.8 像素（96 DPI）
            dpi = 96
            px_per_cm = dpi / 2.54

            # 转换为厘米
            orig_width_cm = orig_width_px / px_per_cm
            orig_height_cm = orig_height_px / px_per_cm

            # 计算缩放比例
            scale_w = max_width_cm / orig_width_cm if orig_width_cm > max_width_cm else 1.0
            scale_h = max_height_cm / orig_height_cm if orig_height_cm > max_height_cm else 1.0

            # 使用较小的缩放比例，确保图片不会超出边界
            scale = min(scale_w, scale_h)

            final_width_cm = orig_width_cm * scale
            final_height_cm = orig_height_cm * scale

            return final_width_cm, final_height_cm

    except Exception as e:
        print(f"[警告] 无法读取图片尺寸: {e}")
        return max_width_cm, None


def add_image(doc, image_path: str, width_cm: float = None, max_width_cm: float = 14.0, max_height_cm: float = 12.0, base_dir: str = ''):
    """
    添加图片到文档（自动缩放）

    修复说明：python-docx 中图片必须通过 doc.add_picture() 添加，
    它会自动创建一个包含图片的新段落。我们再设置该段落居中对齐。

    Args:
        doc: Word 文档对象
        image_path: 图片路径（相对路径或绝对路径）
        width_cm: 图片宽度（厘米），如果为 None 则自动计算
        max_width_cm: 最大宽度限制（厘米）
        max_height_cm: 最大高度限制（厘米）
        base_dir: 基础目录，用于解析相对路径

    Returns:
        bool: 是否成功添加图片
    """
    from pathlib import Path

    # 处理路径
    if base_dir and not Path(image_path).is_absolute():
        full_path = Path(base_dir) / image_path
    else:
        full_path = Path(image_path)

    # 规范化路径（处理 Windows 反斜杠问题）
    full_path = full_path.resolve()

    # 检查文件是否存在
    if not full_path.exists():
        print(f"[警告] 图片文件不存在: {full_path}")
        return False

    # 检查文件扩展名是否为支持的图片格式
    supported_formats = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.tif', '.emf', '.wmf'}
    if full_path.suffix.lower() not in supported_formats:
        print(f"[警告] 不支持的图片格式: {full_path.suffix} (文件: {full_path})")
        return False

    try:
        # 计算合适的尺寸
        if width_cm is None:
            calc_width, calc_height = calculate_image_size(str(full_path), max_width_cm, max_height_cm)
            width_cm = calc_width
        else:
            width_cm = min(width_cm, max_width_cm)

        # 关键修复：使用 doc.add_picture() 直接添加图片
        # 这会自动创建一个包含图片的新段落
        # 之前的 run.add_picture() 不是 python-docx 的标准用法，导致图片不显示
        doc.add_picture(str(full_path), width=Cm(width_cm))

        # 获取刚添加的图片段落（最后一个段落），设置居中对齐
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        last_paragraph.paragraph_format.first_line_indent = Cm(0)
        last_paragraph.paragraph_format.space_before = Pt(6)
        last_paragraph.paragraph_format.space_after = Pt(3)

        print(f"[成功] 图片已插入: {full_path.name} (宽度: {width_cm:.1f}cm)")
        return True
    except Exception as e:
        print(f"[警告] 插入图片失败: {full_path} - 错误: {e}")

        # 失败时在文档中插入占位文字，方便用户手动补图
        try:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(f"[图片缺失: {full_path.name}]")
            set_chinese_font(run, '宋体', 10.5)
            run.font.color.rgb = RGBColor(255, 0, 0)  # 红色标记
        except:
            pass

        return False


def add_figure_caption(doc, caption: str):
    """
    添加图注（图片说明）

    格式：居中、宋体、五号（10.5pt）

    Args:
        doc: Word 文档对象
        caption: 图注文本，如 "图4-1 系统架构图"
    """
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.first_line_indent = Cm(0)
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after = Pt(12)

    run = para.add_run(caption)
    set_chinese_font(run, '宋体', 10.5)  # 五号字体


def add_page_break(doc):
    """
    添加分页符

    在当前位置插入分页符，后续内容从新页面开始。

    Args:
        doc: Word 文档对象
    """
    # 使用 docx 的 add_page_break 方法
    doc.add_page_break()


def convert_md_to_docx(input_path: str, output_path: str) -> Tuple[bool, str]:
    """
    将 Markdown 转换为 Word 文档

    Returns:
        (success, message)
    """
    if not DOCX_AVAILABLE:
        return False, "python-docx 未安装"

    try:
        print(f"[信息] 正在读取: {input_path}")

        # 获取输入文件所在目录，用于解析图片相对路径
        input_file = Path(input_path)
        base_dir = input_file.parent

        preflight_ok, preflight_message = preflight_validate_images(input_file)
        if not preflight_ok:
            return False, preflight_message

        with open(input_path, 'r', encoding='utf-8') as f:
            content = f.read()

        # 导出 docx 时清除 DOI 链接（Markdown 终稿中保留 DOI 便于溯源）
        content = strip_doi_links(content)

        # 清理无意义字符
        content = clean_markdown_content(content)

        # 解析 Markdown
        elements = parse_markdown(content)

        # 创建文档
        doc = create_thesis_document()

        # 统计信息
        image_count = 0
        image_failed = []

        # 处理各元素
        for elem in elements:
            elem_type = elem[0]

            if elem_type == 'title':
                add_title(doc, elem[1])
            elif elem_type == 'h1':
                add_heading(doc, elem[1], 1)
            elif elem_type == 'h2':
                add_heading(doc, elem[1], 2)
            elif elem_type == 'h3':
                add_heading(doc, elem[1], 3)
            elif elem_type == 'para':
                add_paragraph(doc, elem[1])
            elif elem_type == 'code':
                add_code_block(doc, elem[1], elem[2] if len(elem) > 2 else '')
            elif elem_type == 'table':
                add_table(doc, elem[1])
            elif elem_type == 'list':
                add_list_item(doc, elem[1], elem[2] if len(elem) > 2 else False)
            elif elem_type == 'pagebreak':
                # 分页符处理
                add_page_break(doc)
            elif elem_type == 'image':
                # 处理图片（自动缩放）
                img_path = elem[1]
                alt_text = elem[2] if len(elem) > 2 else ''
                success = add_image(doc, img_path, width_cm=None, base_dir=str(base_dir))
                if success:
                    image_count += 1
                    # 添加图注（如果有说明文字）
                    if alt_text:
                        add_figure_caption(doc, alt_text)
                else:
                    image_failed.append(img_path)

        # 保存文档
        doc.save(output_path)
        print(f"[成功] Word 文档已保存: {output_path}")

        # 输出图片统计
        if image_count > 0:
            print(f"[信息] 成功插入 {image_count} 张图片")
        if image_failed:
            print(f"[警告] {len(image_failed)} 张图片插入失败: {image_failed}")

        return True, f"Word 文档已保存到 {output_path}"

    except Exception as e:
        return False, f"转换失败: {str(e)}"
